"""Generate BLUEPRINT.docx from BLUEPRINT.md for Microsoft Word."""

from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    raise SystemExit("Install: pip install python-docx")

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "BLUEPRINT.md"
DOCX_PATH = ROOT / "BLUEPRINT.docx"


def add_table_from_md(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            if j < cols:
                table.rows[i].cells[j].text = cell.strip()


def main() -> None:
    text = MD_PATH.read_text(encoding="utf-8")
    doc = Document()

    # Title page
    title = doc.add_heading("BLUEPRINT PROYEK", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Calculator 2026 Pro")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
    doc.add_paragraph("Dokumen Blueprint Resmi — Versi 1.0")
    doc.add_page_break()

    table_buffer: list[list[str]] = []
    in_code = False
    code_lines: list[str] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            add_table_from_md(doc, table_buffer)
            table_buffer = []

    def flush_code() -> None:
        nonlocal code_lines
        if code_lines:
            p = doc.add_paragraph("\n".join(code_lines))
            for run in p.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            code_lines = []

    for raw in text.splitlines():
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("|") and "|" in line[1:]:
            if set(line.replace("|", "").replace(" ", "")) <= {"-", ":"}:
                continue
            cells = [c.strip() for c in line.strip("|").split("|")]
            table_buffer.append(cells)
            continue
        else:
            flush_table()

        if not line.strip():
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), 0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), 1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), 2)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line[0].isdigit() and ". " in line[:4]:
            doc.add_paragraph(line.split(". ", 1)[1].strip(), style="List Number")
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*"))
            run.bold = True
        else:
            doc.add_paragraph(line)

    flush_table()
    flush_code()

    doc.save(DOCX_PATH)
    print(f"Created: {DOCX_PATH}")


if __name__ == "__main__":
    main()
